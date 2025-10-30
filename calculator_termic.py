import pandas as pd
import streamlit as st
import uuid
import math
import datetime
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt

# --- CONFIGURAREA PAGINII ---
st.set_page_config(page_title="Calculator Termic & Radiatoare PRO", page_icon="️", layout="wide")

# --- INITIALIZAREA st.session_state (Memoria aplicatiei) ---
if 'elemente_anvelopa_salvate' not in st.session_state: st.session_state.elemente_anvelopa_salvate = {"Perete Exterior BCA (exemplu)": 0.28, "Planseu peste subsol (exemplu)": 0.35}
if 'elemente_vitrate_salvate' not in st.session_state: st.session_state.elemente_vitrate_salvate = {"Fereastra Tripan (PVC)": 1.1, "Fereastra Bipan (PVC)": 1.8, "Usa exterioara (izolata)": 1.6}
if 'straturi' not in st.session_state: st.session_state.straturi = []
if 'proiect' not in st.session_state: st.session_state.proiect = {}
if 'elemente_curente' not in st.session_state: st.session_state.elemente_curente = []
if 'punti_curente' not in st.session_state: st.session_state.punti_curente = []
if 'locatie_selectata' not in st.session_state: st.session_state.locatie_selectata = "București (II)"
if 'temp_ext' not in st.session_state: st.session_state.temp_ext = -15
if 'furniture' not in st.session_state: st.session_state.furniture = {} # Pentru modulul IPAT

# --- BAZE DE DATE COMPLETE (VERSIUNE EXTINSĂ) ---
materiale_constructii = {
    # Zidarii
    "BCA (350 kg/m³)": {"lambda": 0.11}, "BCA (500 kg/m³)": {"lambda": 0.14},
    "Caramida Porotherm N+F": {"lambda": 0.26}, "Caramida plina": {"lambda": 0.70}, "Caramida eficienta (BKS)": {"lambda": 0.35},
    # Betoane
    "Beton armat (2400 kg/m³)": {"lambda": 1.75}, "Beton simplu (2200 kg/m³)": {"lambda": 1.40}, "Beton celular usor": {"lambda": 0.20},
    # Termoizolatii
    "Polistiren expandat (EPS 80)": {"lambda": 0.040}, "Polistiren grafitat (EPS 80F)": {"lambda": 0.032},
    "Polistiren extrudat (XPS)": {"lambda": 0.035}, "Vata minerala bazaltica": {"lambda": 0.038}, "Vata minerala de sticla": {"lambda": 0.042},
    "Spuma poliuretanica (PIR)": {"lambda": 0.022}, "Spuma poliuretanica (PUR)": {"lambda": 0.028}, "Pluta expandata": {"lambda": 0.045},
    # Finisaje si diverse
    "Tencuiala ciment-var": {"lambda": 0.87}, "Tencuiala decorativa": {"lambda": 0.70}, "Gips-carton": {"lambda": 0.25},
    "Sapa ciment": {"lambda": 1.40}, "Parchet laminat": {"lambda": 0.17}, "Lemn (ex. OSB)": {"lambda": 0.13}, "Strat de aer neventilat": {"lambda": 0.16},
}
stratificatii_uzuale = {
    "Perete Exterior BCA": [{'material': 'Tencuiala ciment-var', 'grosime_cm': 2.0}, {'material': 'BCA (350 kg/m³)', 'grosime_cm': 30.0}, {'material': 'Polistiren grafitat (EPS 80F)', 'grosime_cm': 15.0}, {'material': 'Tencuiala decorativa', 'grosime_cm': 0.5}],
    "Perete Exterior Cărămidă": [{'material': 'Tencuiala ciment-var', 'grosime_cm': 2.0}, {'material': 'Caramida Porotherm N+F', 'grosime_cm': 25.0}, {'material': 'Vata minerala bazaltica', 'grosime_cm': 15.0}, {'material': 'Tencuiala decorativa', 'grosime_cm': 0.5}],
    "Acoperiș Terasă": [{'material': 'Gips-carton', 'grosime_cm': 1.25}, {'material': 'Beton armat (2400 kg/m³)', 'grosime_cm': 20.0}, {'material': 'Polistiren extrudat (XPS)', 'grosime_cm': 20.0}]
}
radiatoare_otel_catalog = {
    600: {
        "Tip 22": {400: 780, 600: 1170, 800: 1560, 1000: 1950, 1200: 2340, 1400: 2730, 1600: 3120, 1800: 3510, 2000: 3900},
        "Tip 11": {400: 450, 600: 670, 800: 890, 1000: 1120, 1200: 1340, 1400: 1560, 1600: 1780, 1800: 2010, 2000: 2230},
        "Tip 33": {400: 1150, 600: 1720, 800: 2290, 1000: 2870, 1200: 3440, 1400: 4010, 1600: 4590, 1800: 5160, 2000: 5730}
    },
    500: {
        "Tip 22": {400: 680, 600: 1020, 800: 1360, 1000: 1700, 1200: 2040, 1400: 2380, 1600: 2720, 1800: 3060, 2000: 3400}
    },
    900: {
        "Tip 22": {400: 1100, 600: 1650, 800: 2200, 1000: 2750, 1200: 3300, 1400: 3850, 1600: 4400, 1800: 4950, 2000: 5500}
    }
}
rezistente_termice_finisaj = {
    "Gresie / Piatră naturală (<1cm)": {"R": 0.02, "recomandare": "✅ Ideal. Conductivitate termică excelentă."},
    "Parchet laminat specific IPAT (8mm)": {"R": 0.08, "recomandare": "✅ Bun. Alegeți produse cu marcaj specific pentru încălzire în pardoseală."},
    "Parchet triplustratificat (14mm)": {"R": 0.12, "recomandare": "⚠️ Acceptabil. Reduce eficiența sistemului. Necesită temperaturi mai mari ale agentului termic."},
    "Mochetă subțire fără strat de spumă": {"R": 0.10, "recomandare": "⚠️ Acceptabil. Alegeți mochete cu rezistență termică declarată redusă."},
    "Pardoseli LVT / Vinil (click)": {"R": 0.04, "recomandare": "✅ Foarte bun. Material subțire cu transfer termic eficient."},
    "Parchet masiv de stejar (20mm)": {"R": 0.17, "recomandare": "❌ Nerecomandat. Lemnul masiv este un izolator puternic și poate fi deteriorat de variațiile de temperatură."}
}
debite_aer_normate = {"bucatarie": 1.19, "baie": 1, "camara": 1, "camera tehnica": 0.79, "debara": 1, "dormitor": 0.792, "hol": 0.792, "living": 0.792, "sufragerie": 0.792, "birou": 0.792, "default": 0.5}
spatii_adiacente = {"Exterior": {"b": 1.0}, "Sol": {"b": 0.5}, "Subsol neincalzit": {"b": 0.5}, "Pod neincalzit": {"b": 0.7}, "Camera adiacenta incalzita": {"b": 0.0}}
temperaturi_standard = {"dormitor": 20, "living": 20, "baie": 22, "bucatarie": 18, "hol": 18, "default": 20}
zone_climatice = {
    "Manual": -15, "București (II)": -15, "Constanța (I)": -12, "Iași (III)": -18, "Cluj-Napoca (III)": -18, "Timișoara (II)": -15, "Brașov (IV)": -21, "Craiova (II)": -15,
    "Galați (II)": -15, "Ploiești (II)": -15, "Oradea (II)": -15, "Brăila (II)": -15, "Arad (II)": -15, "Pitești (III)": -18, "Sibiu (III)": -18, "Bacău (III)": -18,
    "Târgu Mureș (III)": -18, "Baia Mare (III)": -18, "Buzău (II)": -15, "Botoșani (III)": -18, "Satu Mare (III)": -18, "Râmnicu Vâlcea (III)": -18, "Suceava (III)": -18,
    "Piatra Neamț (III)": -18, "Drobeta-Turnu Severin (II)": -15, "Focșani (II)": -15, "Târgu Jiu (III)": -18, "Tulcea (I)": -12, "Târgoviște (III)": -18, "Slatina (II)": -15,
    "Hunedoara (III)": -18, "Zalău (III)": -18, "Sfântu Gheorghe (IV)": -21, "Vaslui (III)": -18, "Călărași (II)": -15, "Alba Iulia (III)": -18, "Giurgiu (II)": -15, "Miercurea Ciuc (IV)": -21
}
coeficienti_montaj_radiator = {"Montare liberă pe perete": 1.0, "Cu poliță deasupra": 1.05, "În nișă de perete": 1.10, "În mască integrală": 1.25}


# --- FUNCTII UTILITARE ---
def calculeaza_u(straturi):
    R_total = 0.13 + 0.04
    for strat in straturi: R_total += (strat["grosime_cm"] / 100.0) / materiale_constructii.get(strat["material"], {"lambda": 999})["lambda"]
    return 1 / R_total if R_total > 0 else 0

def incarca_stratificatie(nume): st.session_state.straturi = stratificatii_uzuale[nume]

def genereaza_word(data_proiect, temp_ext, locatie):
    document = Document()
    style = document.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    document.add_heading('MEMORIU TEHNIC - INSTALAȚII TERMICE', level=1)
    beneficiar = st.session_state.get('beneficiar', '_________________________'); proiectant = st.session_state.get('proiectant', '_________________________'); data_azi = datetime.date.today().strftime("%d.%m.%Y")
    document.add_heading('1. DATE GENERALE', level=2); document.add_paragraph(f'Proiect: Calculul necesarului de căldură\nBeneficiar: {beneficiar}\nAmplasament: {locatie}\nData: {data_azi}')
    document.add_heading('2. BAZA DE CALCUL', level=2); document.add_paragraph(f'Normativ de referință: SR 12831\nTemperatură exterioară de calcul: {temp_ext}°C (zona climatică {locatie}).')
    document.add_heading('3. REZULTATELE CALCULULUI', level=2)
    df_rezumat = pd.DataFrame([{'Încăpere': nume, 'Transmisie [W]': date['Q_Trans_Total'], 'Ventilație [W]': date['Q_Vent'], 'Necesar Total [W]': date['Q_Total']} for nume, date in data_proiect.items()])
    total_cladire = df_rezumat['Necesar Total [W]'].sum()
    table = document.add_table(rows=1, cols=4); table.style = 'Table Grid'; hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Încăpere'; hdr_cells[1].text = 'Transmisie [W]'; hdr_cells[2].text = 'Ventilație [W]'; hdr_cells[3].text = 'Necesar Total [W]'
    for index, row in df_rezumat.iterrows():
        row_cells = table.add_row().cells; row_cells[0].text = row['Încăpere']; row_cells[1].text = f"{row['Transmisie [W]']:.0f}"; row_cells[2].text = f"{row['Ventilație [W]']:.0f}"; row_cells[3].text = f"**{row['Necesar Total [W]']:.0f}**"
    total_cells = table.add_row().cells; total_cells[0].merge(total_cells[2]); total_cells[0].text = 'TOTAL CLĂDIRE'; total_cells[3].text = f"**{total_cladire:.0f} W**"
    document.add_heading('4. CONCLUZII', level=2); document.add_paragraph(f'Necesarul total de căldură pentru clădire este de {total_cladire:.0f} W.\nSursa de căldură (centrală termică, pompă de căldură etc.) va trebui să aibă o putere nominală cel puțin egală cu această valoare.')
    document.add_paragraph(f'\n\nÎntocmit,\n{proiectant}')
    bio = io.BytesIO(); document.save(bio); return bio.getvalue()

def genereaza_pdf(data_proiect, temp_ext, locatie):
    # ATENTIE: Aceasta functie necesita ca fisierul DejaVuSans.ttf sa fie in acelasi folder cu scriptul
    # Il puteti descarca de pe Google Fonts.
    pdf = FPDF()
    pdf.add_page()
    try:
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 12)
    except RuntimeError:
        pdf.set_font('Arial', '', 12)
        st.warning("Fontul DejaVu nu a fost găsit. Diacriticele s-ar putea să nu fie afișate corect. Descărcați 'DejaVuSans.ttf' și plasați-l lângă script.", icon="⚠️")

    beneficiar = st.session_state.get('beneficiar', 'Nespecificat'); proiectant = st.session_state.get('proiectant', 'Nespecificat'); data_azi = datetime.date.today().strftime("%d.%m.%Y")
    pdf.set_font_size(16); pdf.cell(0, 10, 'MEMORIU TEHNIC - INSTALAȚII TERMICE', 0, 1, 'C'); pdf.ln(10)
    pdf.set_font_size(12); pdf.multi_cell(0, 8, f"**1. DATE GENERALE**\n- **Proiect:** Calculul necesarului de căldură\n- **Beneficiar:** {beneficiar}\n- **Amplasament:** {locatie}\n- **Data:** {data_azi}", markdown=True); pdf.ln(5)
    pdf.multi_cell(0, 8, f"**2. BAZA DE CALCUL**\n- **Normativ de referință:** SR 12831\n- **Temperatură exterioară de calcul:** {temp_ext}°C (zona climatică {locatie}).", markdown=True); pdf.ln(5)
    pdf.set_font_size(14); pdf.cell(0, 10, '3. REZULTATE CALCUL NECESAR DE CĂLDURĂ', 0, 1, 'L'); pdf.set_font_size(10)
    line_height = pdf.font_size * 2; col_width = [60, 40, 40, 40]
    pdf.cell(col_width[0], line_height, 'Încăpere', border=1); pdf.cell(col_width[1], line_height, 'Transmisie [W]', border=1); pdf.cell(col_width[2], line_height, 'Ventilatie [W]', border=1); pdf.cell(col_width[3], line_height, 'Necesar Total [W]', border=1); pdf.ln(line_height)
    total_cladire = 0
    for nume, date in data_proiect.items():
        pdf.cell(col_width[0], line_height, nume, border=1); pdf.cell(col_width[1], line_height, f"{date['Q_Trans_Total']:.0f}", border=1); pdf.cell(col_width[2], line_height, f"{date['Q_Vent']:.0f}", border=1); pdf.cell(col_width[3], line_height, f"{date['Q_Total']:.0f}", border=1); pdf.ln(line_height)
        total_cladire += date['Q_Total']
    pdf.set_font(style='B'); pdf.cell(col_width[0] + col_width[1] + col_width[2], line_height, 'TOTAL CLĂDIRE', border=1); pdf.cell(col_width[3], line_height, f"{total_cladire:.0f} W", border=1); pdf.ln(15)
    pdf.set_font_size(12); pdf.set_font(style=''); pdf.multi_cell(0, 8, f"**4. CONCLUZII**\nNecesarul total de căldură pentru clădire este de **{total_cladire:.0f} W**.", markdown=True); pdf.ln(20)
    pdf.cell(0, 10, f"Întocmit,\n{proiectant}")
    return pdf.output(dest='S').encode('latin-1')

# --- INTERFATA STREAMLIT ---
st.title("️ Calculator Profesional: Necesar Termic & Sisteme de Încălzire")
tab_titles = ["1. 🏛️ Bibliotecă Elemente", "2. 🏗️ Calcul Încăpere", "3. 📊 Centralizatoare", "4. 🌡️ Radiatoare", "5. 🔥 Încălzire Pardoseală", "6. 📝 Generare Memoriu"]
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(tab_titles)

with tab1: # BIBLIOTECA DE ELEMENTE
    st.header("Creează și Gestionează Biblioteca de Elemente")
    col1, col2 = st.columns([1,1])
    with col1:
        with st.container(border=True):
            st.subheader("Creează Element Nou")
            nume_element_nou = st.text_input("Nume element (ex: Perete Exterior, Fereastră Tripan)")
            is_vitrat = st.checkbox("Bifează dacă este element vitrat (fereastră, ușă)")
            if is_vitrat:
                u_val_manual = st.number_input("Introdu valoarea U [W/m²K] direct", value=1.1, format="%.3f")
                if st.button("💾 Salvează Element Vitrat", type="primary", use_container_width=True):
                    if nume_element_nou: st.session_state.elemente_vitrate_salvate[nume_element_nou] = u_val_manual; st.success(f"Elementul '{nume_element_nou}' a fost salvat!"); st.rerun()
                    else: st.warning("Introdu un nume.")
            else:
                with st.expander("Încarcă Stratificații Uzuale (Exemple)"):
                    c1,c2,c3=st.columns(3); c1.button("Perete BCA", on_click=incarca_stratificatie, args=("Perete Exterior BCA",), use_container_width=True); c2.button("Perete Cărămidă", on_click=incarca_stratificatie, args=("Perete Exterior Cărămidă",), use_container_width=True); c3.button("Terasă", on_click=incarca_stratificatie, args=("Acoperiș Terasă",), use_container_width=True)
                if st.button("➕ Adaugă Strat Nou", use_container_width=True): st.session_state.straturi.append({'material': 'Polistiren expandat (EPS 80)', 'grosime_cm': 10.0})
                for i, strat in enumerate(st.session_state.straturi):
                    cols = st.columns([3, 2, 0.5]); st.session_state.straturi[i]['material'] = cols[0].selectbox(f"Material Strat {i+1}", list(materiale_constructii.keys()), key=f"mat_{i}", index=max(0, list(materiale_constructii.keys()).index(strat['material']) if strat.get('material') in materiale_constructii else 0))
                    st.session_state.straturi[i]['grosime_cm'] = cols[1].number_input(f"Grosime [cm]", min_value=0.1, value=float(strat.get('grosime_cm', 10.0)), key=f"gros_{i}")
                    if cols[2].button("🗑️", key=f"del_strat_{i}", help="Șterge strat"): st.session_state.straturi.pop(i); st.rerun()
                if st.session_state.straturi:
                    U_val = calculeaza_u(st.session_state.straturi)
                    st.metric(label="Coeficient U calculat", value=f"{U_val:.3f} W/m²K")
                    if st.button("💾 Salvează Element de Anvelopă", type="primary", use_container_width=True):
                        if nume_element_nou: st.session_state.elemente_anvelopa_salvate[nume_element_nou] = U_val; st.success(f"Elementul '{nume_element_nou}' a fost salvat!"); st.rerun()
                        else: st.warning("Introdu un nume.")
    with col2:
        with st.container(border=True):
            st.subheader("Elemente de Anvelopă Salvate"); [st.code(f"{nume}: U = {u_val:.3f} W/m²K") for nume, u_val in st.session_state.elemente_anvelopa_salvate.items()]
        with st.container(border=True):
            st.subheader("Elemente Vitrate Salvate"); [st.code(f"{nume}: U = {u_val:.3f} W/m²K") for nume, u_val in st.session_state.elemente_vitrate_salvate.items()]

with tab2: # CALCUL PE INCAPERE
    st.header("Definește o Încăpere și Calculează Pierderile")
    with st.container(border=True):
        st.subheader("📍 1. Setări Generale Proiect")
        locatie_selectata = st.selectbox("Selectează Zona Climatică", options=list(zone_climatice.keys()), index=list(zone_climatice.keys()).index(st.session_state.locatie_selectata))
        st.session_state.locatie_selectata = locatie_selectata; temp_ext = zone_climatice[locatie_selectata]; st.session_state.temp_ext = temp_ext
        st.info(f"Temperatura exterioară de calcul pentru **{locatie_selectata}** este: **{temp_ext}°C**")

    with st.container(border=True):
        st.subheader("🏠 2. Date Generale Încăpere")
        c1, c2, c3, c4 = st.columns(4)
        nume_incapere = c1.text_input("Nume Încăpere", f"Incăpere {len(st.session_state.proiect) + 1}")
        temp_int = c2.number_input("Temp. interioară [°C]", value=20.0)
        lungime = c3.number_input("Lungime [m]", value=4.0, min_value=0.1); latime = c4.number_input("Lățime [m]", value=3.5, min_value=0.1)
        inaltime = st.number_input("Înălțime [m]", value=2.7, min_value=0.1)
        tip_incapere = next((k for k in temperaturi_standard if k in nume_incapere.lower()), "default")
        if abs(temp_int - temperaturi_standard[tip_incapere]) > 2: st.warning(f"⚠️ Atenție: Temperatura standard pentru '{tip_incapere.capitalize()}' este de {temperaturi_standard[tip_incapere]}°C.")

    st.subheader("🧱 3. Elemente de Anvelopă (Pereți, Planșee)")
    tip_adauga = st.selectbox("Asistent Adăugare Elemente", ["Adaugă Pardoseală/Tavan (L x l)", "Adaugă Perete (pe Lungime L x H)", "Adaugă Perete (pe lățime l x H)", "Adaugă Element Manual"])
    if st.button("➕ Adaugă Element Anvelopă", use_container_width=True):
        arie_sugerata = 10.0
        if "Pardoseală" in tip_adauga: arie_sugerata = lungime * latime
        elif "Lungime" in tip_adauga: arie_sugerata = lungime * inaltime
        elif "lățime" in tip_adauga: arie_sugerata = latime * inaltime
        st.session_state.elemente_curente.append({'id': uuid.uuid4(), 'tip_el': list(st.session_state.elemente_anvelopa_salvate.keys())[0], 'spatiu_adj': 'Exterior', 'arie_tot': arie_sugerata, 'vitraje': [], 'temp_adj_dif': False, 'temp_adj_val': 18.0})

    for el_idx, el in enumerate(st.session_state.elemente_curente):
        with st.container(border=True):
            c1,c2,c3, c4 = st.columns([2,1,1, 0.5]); el['tip_el'] = c1.selectbox("Element", list(st.session_state.elemente_anvelopa_salvate.keys()), key=f"sel_{el['id']}"); el['spatiu_adj'] = c2.selectbox("Spațiu Adiacent", list(spatii_adiacente.keys()), key=f"spa_{el['id']}"); el['arie_tot'] = c3.number_input("Arie Totală [m²]", value=el['arie_tot'], min_value=0.1, key=f"ar_{el['id']}");
            if c4.button("🗑️", key=f"del_el_{el['id']}", help="Șterge element"): st.session_state.elemente_curente.pop(el_idx); st.rerun()
            if el['spatiu_adj'] == "Camera adiacenta incalzita":
                el['temp_adj_dif'] = st.checkbox("Temperatură diferită?", key=f"check_{el['id']}", value=el['temp_adj_dif'])
                if el['temp_adj_dif']: el['temp_adj_val'] = st.number_input("Temp. adiacentă [°C]", value=el['temp_adj_val'], key=f"temp_adj_{el['id']}")

            with st.expander("Adaugă Ferestre/Uși în acest element"):
                if st.button("➕ Vitraj", key=f"add_v_{el['id']}"): el['vitraje'].append({'id': uuid.uuid4(), 'tip_v': list(st.session_state.elemente_vitrate_salvate.keys())[0], 'L': 1.2, 'H': 1.4})
                for i, vitraj in enumerate(el['vitraje']):
                    cv1,cv2,cv3,cv4 = st.columns([2, 1, 1, 0.5]); vitraj['tip_v']=cv1.selectbox("Tip Vitraj", list(st.session_state.elemente_vitrate_salvate.keys()), key=f"v_tip_{vitraj['id']}"); vitraj['L']=cv2.number_input("Lățime [m]", value=vitraj['L'], min_value=0.1, key=f"v_l_{vitraj['id']}"); vitraj['H']=cv3.number_input("Înălțime [m]", value=vitraj['H'], min_value=0.1, key=f"v_h_{vitraj['id']}");
                    if cv4.button("🗑️", key=f"del_v_{vitraj['id']}", help="Șterge vitraj"): el['vitraje'].pop(i); st.rerun()

            arie_vitraje = sum(v['L']*v['H'] for v in el.get('vitraje', [])); arie_neta = el['arie_tot'] - arie_vitraje
            if arie_neta < 0: st.error(f"Suprafața vitrajelor ({arie_vitraje:.2f} m²) depășește suprafața totală a elementului ({el['arie_tot']:.2f} m²).")
            st.metric(label="Suprafață Opacă Netă (de calcul)", value=f"{arie_neta:.2f} m²")


    st.subheader("🧊 4. Punți Termice")
    if st.button("➕ Adaugă Punte Termică", use_container_width=True):
        st.session_state.punti_curente.append({'id': uuid.uuid4(), 'descr': 'Contur fereastră', 'lungime': 5.2, 'psi': 0.04})
    for pt_idx, pt in enumerate(st.session_state.punti_curente):
        with st.container(border=True):
            c1,c2,c3,c4 = st.columns([3,1,1,0.5]); pt['descr'] = c1.text_input("Descriere", value=pt['descr'], key=f"pt_d_{pt['id']}"); pt['lungime'] = c2.number_input("Lungime [m]", value=pt['lungime'], min_value=0.1, key=f"pt_l_{pt['id']}"); pt['psi'] = c3.number_input("Valoare Psi [W/mK]", value=pt['psi'], format="%.3f", help="Ex: Colț=0.05", key=f"pt_psi_{pt['id']}");
            if c4.button("🗑️", key=f"del_pt_{pt['id']}", help="Șterge punte termică"): st.session_state.punti_curente.pop(pt_idx); st.rerun()

    st.subheader("🏁 5. Finalizare Calcul Încăpere")
    if st.button("✅ Calculează și Adaugă Încăperea în Proiect", type="primary", use_container_width=True):
        tip_incapere_vent = next((k for k in debite_aer_normate if k in nume_incapere.lower()), "default"); V = lungime*latime*inaltime; n = debite_aer_normate[tip_incapere_vent]
        Q_vent = 0.34 * V * n * (temp_int - temp_ext)
        Q_trans_anvelopa = 0; centralizator = []
        for el in st.session_state.elemente_curente:
            b = spatii_adiacente[el['spatiu_adj']]['b']; delta_T = temp_int - temp_ext
            if el['spatiu_adj'] == "Camera adiacenta incalzita":
                if el['temp_adj_dif']: delta_T = temp_int - el['temp_adj_val']; b = 1.0
                else: delta_T = 0
            arie_vitraje = sum(v['L']*v['H'] for v in el['vitraje']); arie_neta = el['arie_tot'] - arie_vitraje; U = st.session_state.elemente_anvelopa_salvate[el['tip_el']]
            Q_el = b * U * arie_neta * delta_T; Q_trans_anvelopa += Q_el
            if arie_neta > 0.01: centralizator.append({'Element': el['tip_el'], 'Arie [m²]': arie_neta, 'U': U, 'b': b, 'ΔT': delta_T, 'Q_Transmisie [W]': Q_el})
            for v in el['vitraje']:
                U_v = st.session_state.elemente_vitrate_salvate[v['tip_v']]; arie_v = v['L']*v['H']
                Q_v = b * U_v * arie_v * delta_T; Q_trans_anvelopa += Q_v
                centralizator.append({'Element': v['tip_v'], 'Arie [m²]': arie_v, 'U': U_v, 'b': b, 'ΔT': delta_T, 'Q_Transmisie [W]': Q_v})
        Q_punti = sum(pt['lungime'] * pt['psi'] * (temp_int - temp_ext) for pt in st.session_state.punti_curente)
        Q_trans_total = Q_trans_anvelopa + Q_punti; Q_total = Q_trans_total + Q_vent
        st.session_state.proiect[nume_incapere] = {'Q_Vent': Q_vent, 'Q_Trans_Anvelopa': Q_trans_anvelopa, 'Q_Punti': Q_punti, 'Q_Trans_Total': Q_trans_total, 'Q_Total': Q_total, 'centralizator': centralizator, 'temp_int': temp_int, 'lungime': lungime, 'latime': latime}
        st.session_state.elemente_curente, st.session_state.punti_curente = [], []; st.success(f"Încăperea '{nume_incapere}' a fost adăugată cu un necesar total de {Q_total:.0f} W."); st.rerun()

with tab3: # CENTRALIZATOARE
    st.header("Vizualizare Rezultate Proiect")
    if not st.session_state.proiect: st.warning("Niciun calcul efectuat. Adaugă o încăpere din Tab-ul 2.")
    else:
        st.subheader("Rezumat pe Încăperi")
        rezumat_data = [{'Încăpere': nume, 'Transmisie Anvelopă [W]': date['Q_Trans_Anvelopa'], 'Punți Termice [W]': date['Q_Punti'], 'Total Transmisie [W]': date['Q_Trans_Total'], 'Ventilație [W]': date['Q_Vent'], 'Necesar Total [W]': date['Q_Total']} for nume, date in st.session_state.proiect.items()]
        df_rezumat = pd.DataFrame(rezumat_data)
        coloane_format = [col for col in df_rezumat.columns if col != 'Încăpere']
        st.dataframe(df_rezumat.style.format("{:.0f}", subset=coloane_format))
        st.metric("Necesar Total Clădire", f"{df_rezumat['Necesar Total [W]'].sum():.0f} W")
        with st.expander("Vezi Centralizatorul Detaliat (pentru Verificare)"):
            full_centralizator = [];
            for nume, date in st.session_state.proiect.items():
                for row in date['centralizator']: new_row = row.copy(); new_row['Încăpere'] = nume; full_centralizator.append(new_row)
            df_centralizator = pd.DataFrame(full_centralizator); st.dataframe(df_centralizator.style.format("{:.2f}", subset=['Arie [m²]', 'U', 'b', 'ΔT', 'Q_Transmisie [W]']))
            st.download_button("Descarcă Centralizator Detaliat (.csv)", df_centralizator.to_csv(index=False).encode('utf-8'), "centralizator_detaliat.csv", "text/csv")

with tab4: # DIMENSIONARE RADIATOARE
    st.header("Ghid pentru Dimensionarea Radiatoarelor")
    if not st.session_state.proiect: st.warning("Calculează necesarul termic pentru a putea dimensiona radiatoarele.")
    else:
        with st.container(border=True):
            st.subheader("Pasul 1: Setează Parametrii Sistemului de Încălzire")
            c1, c2 = st.columns(2); temp_tur = c1.number_input("Temperatură Tur [°C]", value=75.0); temp_retur = c2.number_input("Temperatură Retur [°C]", value=65.0)
        st.subheader("Pasul 2: Alege Radiatorul pentru Fiecare Cameră")
        for nume, date in st.session_state.proiect.items():
            with st.container(border=True):
                st.markdown(f"#### {nume} (Necesar: **{date['Q_Total']:.0f} W**)")
                c1,c2,c3,c4 = st.columns(4)
                delta_T_ref = c1.number_input("ΔT de Referință (din catalog)", value=50.0, min_value=1.0, key=f"dt_ref_{nume}", help="Puterea din fișa tehnică este dată pentru un ΔT standard (de obicei 50K sau 60K).")
                mod_montaj = c2.selectbox("Mod de montare", list(coeficienti_montaj_radiator.keys()), key=f"montaj_{nume}", help="Un radiator acoperit pierde din eficiență. Coeficientul compensează.")
                inaltime_radiator = c3.selectbox("Înălțime Radiator [mm]", options=list(radiatoare_otel_catalog.keys()), key=f"inaltime_{nume}")
                tip_radiator = c4.selectbox("Tip Radiator", options=list(radiatoare_otel_catalog[inaltime_radiator].keys()), key=f"tip_{nume}")
                
                temp_medie_agent = (temp_tur + temp_retur) / 2; delta_T_proiect = temp_medie_agent - date['temp_int']; coef_montaj = coeficienti_montaj_radiator[mod_montaj]
                if delta_T_proiect <= 0: delta_T_proiect = 0.1
                putere_necesara_catalog = date['Q_Total'] * math.pow((delta_T_ref / delta_T_proiect), 1.3) * coef_montaj
                st.metric("🎯 Putere Necesară (la ΔT ref din catalog)", f"{putere_necesara_catalog:.0f} W")
                
                radiator_sugerat = "N/A"; putere_radiator_ales = 0
                catalog_ales = radiatoare_otel_catalog[inaltime_radiator][tip_radiator]
                for lungime, putere in sorted(catalog_ales.items()):
                    if putere >= putere_necesara_catalog:
                        radiator_sugerat = f"{tip_radiator} / {inaltime_radiator} / {lungime} mm"; putere_radiator_ales = putere; break
                
                if radiator_sugerat != "N/A": st.success(f"✅ Radiator Sugerat: **{radiator_sugerat}** (oferă {putere_radiator_ales} W)")
                else: st.warning("⚠️ Puterea necesară este prea mare pentru catalogul selectat. Alegeți un radiator mai mare sau adăugați mai multe corpuri.")

with tab5: # ÎNCĂLZIRE PARDOSALĂ
    st.header("Dimensionare Sistem de Încălzire în Pardoseală (IP)")
    if not st.session_state.proiect: st.warning("Calculează necesarul termic pentru a putea dimensiona sistemul IP.")
    else:
        tip_sistem_ip = st.radio("Selectează Tipul Sistemului de Încălzire în Pardoseală", ("Agent Termic (Hidronic)", "Electric"), horizontal=True)
        st.subheader("Dimensionare Circuite pe Încăpere")
        for nume, date in st.session_state.proiect.items():
            if nume not in st.session_state.furniture: st.session_state.furniture[nume] = []
            with st.container(border=True):
                st.markdown(f"#### {nume} (Necesar: **{date['Q_Total']:.0f} W**)")
                with st.expander("Asistent Calcul Arie Activă"):
                    arie_bruta = date.get('lungime', 5) * date.get('latime', 4)
                    st.write(f"Arie brută încăpere: {arie_bruta:.2f} m²")
                    if st.button("➕ Adaugă piesă mobilier", key=f"add_furn_{nume}"): st.session_state.furniture[nume].append({"L": 1.8, "l": 0.6})
                    total_mobila = 0
                    for i, item in enumerate(st.session_state.furniture[nume]):
                        c1,c2,c3 = st.columns([2,2,1]); item['L'] = c1.number_input("L mobila [m]", value=item['L'], key=f"furn_L_{nume}_{i}"); item['l'] = c2.number_input("l mobila [m]", value=item['l'], key=f"furn_l_{nume}_{i}"); total_mobila += item['L'] * item['l']
                        if c3.button("🗑️", key=f"del_furn_{nume}_{i}"): st.session_state.furniture[nume].pop(i); st.rerun()
                    arie_activa_calculata = arie_bruta - total_mobila
                    st.metric("Arie Activă (de calcul) ⬇️", f"{arie_activa_calculata:.2f} m²")

                arie_activa = st.number_input("Suprafață Activă IP [m²]", value=round(arie_activa_calculata, 2), key=f"arie_ip_{nume}")
                
                if tip_sistem_ip == "Agent Termic (Hidronic)":
                    c1,c2 = st.columns(2); tip_finisaj = c1.selectbox("Tip finisaj pardoseală", list(rezistente_termice_finisaj.keys()), key=f"finisaj_ip_{nume}"); info_finisaj = rezistente_termice_finisaj[tip_finisaj]; c2.info(info_finisaj['recomandare'])
                    if arie_activa > 0:
                        q_necesar = date['Q_Total'] / arie_activa; st.metric("Putere Specifică Necesară", f"{q_necesar:.1f} W/m²")
                        if q_necesar > 100: st.error("❌ Putere specifică foarte mare. Încălzirea în pardoseală s-ar putea să nu facă față singură. Considerați un calorifer suplimentar.")
                        else:
                            if q_necesar > 65: pas_recomandat = 10
                            elif q_necesar > 40: pas_recomandat = 15
                            else: pas_recomandat = 20
                            lungime_totala_teava = arie_activa / (pas_recomandat / 100.0); nr_circuite = math.ceil(lungime_totala_teava / 90); lungime_per_circuit = lungime_totala_teava / nr_circuite if nr_circuite > 0 else 0
                            st.success(f"**Soluție Tehnică Recomandată:**"); c_rez1, c_rez2, c_rez3 = st.columns(3); c_rez1.metric("Pas de montaj", f"{pas_recomandat} cm"); c_rez2.metric("Nr. Circuite", f"{nr_circuite}"); c_rez3.metric("Lungime / Circuit", f"~{lungime_per_circuit:.0f} m")
                
                elif tip_sistem_ip == "Electric":
                    putere_specifica_mat = st.selectbox("Putere specifică covor electric", (100, 150, 200), index=1, key=f"mat_pwr_{nume}", help="150 W/m² este standard pentru spații de locuit.")
                    if arie_activa > 0:
                        putere_necesara_camera = date['Q_Total']; arie_necesara_covor = putere_necesara_camera / putere_specifica_mat
                        st.metric("Putere totală necesară", f"{putere_necesara_camera:.0f} W")
                        if arie_necesara_covor > arie_activa: st.error(f"Aria necesară a covorului ({arie_necesara_covor:.2f} m²) depășește aria activă disponibilă. Alegeți un covor cu putere specifică mai mare sau o sursă suplimentară.")
                        else: st.success(f"✅ **Soluție Recomandată:** Alegeți un covor electric de **{putere_specifica_mat} W/m²** cu o suprafață de cel puțin **{arie_necesara_covor:.2f} m²**.")

with tab6: # GENERARE MEMORIU
    st.header("Generare Memoriu Tehnic de Specialitate")
    if not st.session_state.proiect: st.warning("Finalizează calculul pentru a putea genera memoriul.")
    else:
        st.info("Completează datele de mai jos, apoi apasă butonul pentru a genera textul memoriului și a descărca fișierele.")
        c1, c2 = st.columns(2); st.session_state.beneficiar = c1.text_input("Nume Beneficiar", value=st.session_state.get('beneficiar', '')); st.session_state.proiectant = c2.text_input("Nume Proiectant / Întocmit", value=st.session_state.get('proiectant', ''))
        
        if st.button("🚀 Generează Memoriul", type="primary", use_container_width=True):
            df_rezumat_final = pd.DataFrame([{'Încăpere': nume, 'Transmisie Anvelopă [W]': date['Q_Trans_Anvelopa'], 'Punți Termice [W]': date['Q_Punti'], 'Ventilație [W]': date['Q_Vent'], 'Necesar Total [W]': date['Q_Total']} for nume, date in st.session_state.proiect.items()])
            total_cladire = df_rezumat_final['Necesar Total [W]'].sum(); data_azi = datetime.date.today().strftime("%d.%m.%Y"); locatie_memoriu = st.session_state.locatie_selectata; temp_ext_memoriu = st.session_state.temp_ext
            memoriu = f"""## MEMORIU TEHNIC - INSTALAȚII TERMICE\n\n**1. DATE GENERALE**\n- **Proiect:** Calculul necesarului de căldură\n- **Beneficiar:** {st.session_state.beneficiar if st.session_state.beneficiar else "..."}\n- **Amplasament:** {locatie_memoriu}\n- **Data:** {data_azi}\n\n**2. BAZA DE CALCUL**\n- **Normativ:** SR 12831\n- **Temp. exterioară:** {temp_ext_memoriu}°C\n\n**3. REZULTATE**\n| Încăpere | Transmisie [W] | Necesar Total [W] |\n|---|---|---|\n"""
            for index, row in df_rezumat_final.iterrows(): memoriu += f"| {row['Încăpere']} | {row['Transmisie Anvelopă [W]'] + row['Punți Termice [W]']:.0f} | **{row['Necesar Total [W]']:.0f}** |\n"
            memoriu += f"| **TOTAL** | - | **{total_cladire:.0f}** |\n\n**4. CONCLUZII**\nNecesarul total de căldură este de **{total_cladire:.0f} W**.\n\n**Întocmit,**\n{st.session_state.proiectant if st.session_state.proiectant else "..."}"
            st.code(memoriu, language="markdown")
            st.success("Memoriul a fost generat! Folosește butoanele de mai jos pentru a descărca.")

            c1, c2 = st.columns(2)
            pdf_data = genereaza_pdf(st.session_state.proiect, st.session_state.temp_ext, st.session_state.locatie_selectata)
            c1.download_button(label="📥 Descarcă Memoriu (.pdf)", data=pdf_data, file_name=f"Memoriu_Tehnic_{st.session_state.get('beneficiar', 'Proiect')}.pdf", mime="application/pdf", use_container_width=True)
            word_data = genereaza_word(st.session_state.proiect, st.session_state.temp_ext, st.session_state.locatie_selectata)
            c2.download_button(label="📥 Descarcă Memoriu (.docx)", data=word_data, file_name=f"Memoriu_Tehnic_{st.session_state.get('beneficiar', 'Proiect')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
